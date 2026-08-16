from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.models.documents import Segment
from app.utils.errors import api_error

ALIGNMENTS = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


def _run_info(run) -> dict:
    font = run.font

    info: dict = {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "underline": bool(run.underline),
        "strikethrough": run.font.strike,
        "font": font.name,
        "fontSize": font.size.pt if font.size else None,
    }
    if font.color and font.color.rgb:
        info["color"] = str(font.color.rgb)
    if font.highlight_color is not None:
        info["highlight"] = str(font.highlight_color)
    return info


def _is_list_paragraph(paragraph) -> bool:
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def _list_level(paragraph) -> int | None:
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is not None:
        return int(ilvl.get(qn("w:val"), "0"))
    return 0


class DocumentParser:
    def validate(self, content: bytes) -> None:

        if not content.startswith(b"PK"):
            raise api_error(
                422, "invalid_docx", "The uploaded file is not a valid DOCX archive."
            )
        try:
            with ZipFile(BytesIO(content)) as archive:
                if (
                    "word/document.xml" not in archive.namelist()
                    or "[Content_Types].xml" not in archive.namelist()
                ):
                    raise api_error(
                        422,
                        "invalid_docx",
                        "The uploaded file is not a valid Word document.",
                    )
        except BadZipFile as exc:
            raise api_error(
                422, "invalid_docx", "The uploaded file is not a valid DOCX archive."
            ) from exc

    def parse_file(self, path: Path) -> tuple[list[Segment], int]:
        return self.parse_bytes(path.read_bytes())

    def parse_bytes(self, content: bytes) -> tuple[list[Segment], int]:
        self.validate(content)
        document = Document(BytesIO(content))
        segments: list[Segment] = []
        position = 0

        for index, paragraph in enumerate(document.paragraphs):

            style = paragraph.style.name if paragraph.style else ""
            is_list = _is_list_paragraph(paragraph)

            if style.lower().startswith("heading"):
                kind = "heading"
            elif is_list:
                kind = "list"
            else:
                kind = "paragraph"

            runs = [_run_info(run) for run in paragraph.runs]
            pf = paragraph.paragraph_format

            meta: dict = {
                "style": style,
                "alignment": ALIGNMENTS.get(paragraph.alignment),
                "runs": runs,
                "spaceBefore": pf.space_before.pt if pf.space_before else None,
                "spaceAfter": pf.space_after.pt if pf.space_after else None,
                "lineSpacing": (
                    float(pf.line_spacing)
                    if isinstance(pf.line_spacing, (int, float))
                    else None
                ),
                "location": "body",
            }

            if kind == "heading":
                try:
                    meta["level"] = int(style.split()[-1])
                except (ValueError, IndexError):
                    meta["level"] = 1

            if is_list:
                meta["listLevel"] = _list_level(paragraph)

            if pf.first_line_indent:
                meta["firstLineIndent"] = pf.first_line_indent.pt

            segments.append(
                Segment(
                    id=f"word/document.xml:paragraph:{index}",
                    type=kind,
                    text=paragraph.text,
                    position=position,
                    meta=meta,
                )
            )
            position += 1

        for table_index, table in enumerate(document.tables):

            cells = [[cell.text for cell in row.cells] for row in table.rows]
            cell_formatting: list[list[str]] = []
            for row in table.rows:
                row_fmt = []
                for cell in row.cells:
                    fmt_parts = []
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.bold:
                                fmt_parts.append("bold")
                            if run.italic:
                                fmt_parts.append("italic")
                    row_fmt.append(
                        ",".join(sorted(set(fmt_parts))) if fmt_parts else ""
                    )
                cell_formatting.append(row_fmt)

            segments.append(
                Segment(
                    id=f"word/document.xml:table:{table_index}",
                    type="table",
                    text="\n".join(" | ".join(row) for row in cells),
                    position=position,
                    meta={
                        "rows": len(cells),
                        "cols": len(cells[0]) if cells else 0,
                        "cells": cells,
                        "cellFormatting": cell_formatting,
                        "location": "body",
                    },
                )
            )
            position += 1

        seen_parts: set[int] = set()

        for section_index, section in enumerate(document.sections):
            for location, container in (
                ("header", section.header),
                ("footer", section.footer),
            ):
                marker = id(container._element)
                if marker in seen_parts:
                    continue
                seen_parts.add(marker)
                for paragraph_index, paragraph in enumerate(container.paragraphs):
                    if not paragraph.text:
                        continue
                    runs = [_run_info(run) for run in paragraph.runs]
                    segments.append(
                        Segment(
                            id=f"word/{location}{section_index}.xml:paragraph:{paragraph_index}",
                            type="paragraph",
                            text=paragraph.text,
                            position=position,
                            meta={
                                "style": (
                                    paragraph.style.name if paragraph.style else ""
                                ),
                                "runs": runs,
                                "location": location,
                            },
                        )
                    )
                    position += 1

        image_count = 0

        dimensions = [(shape.width, shape.height) for shape in document.inline_shapes]
        with ZipFile(BytesIO(content)) as archive:
            for zip_path in archive.namelist():
                if zip_path.startswith("word/media/"):
                    width, height = (
                        dimensions[image_count]
                        if image_count < len(dimensions)
                        else (None, None)
                    )
                    segments.append(
                        Segment(
                            id=f"word/document.xml:image:{image_count}",
                            type="image",
                            text=Path(zip_path).name,
                            position=position,
                            meta={
                                "imagePath": zip_path,
                                "widthEmu": width,
                                "heightEmu": height,
                                "widthCm": round(width / 360000, 2) if width else None,
                                "heightCm": (
                                    round(height / 360000, 2) if height else None
                                ),
                                "location": "body",
                            },
                        )
                    )
                    position += 1
                    image_count += 1

        return segments, image_count
