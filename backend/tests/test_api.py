from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient


def make_docx() -> bytes:
    document = Document()
    document.add_heading("Pendahuluan", level=1)
    document.add_paragraph("Nama saya HUSNUL IBAT.")
    buffer = BytesIO(); document.save(buffer)
    return buffer.getvalue()


def client(tmp_path: Path, monkeypatch) -> TestClient:
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path / "data"))
    monkeypatch.setenv("GEMINI_API_KEY", "")
    from app.config import get_settings
    get_settings.cache_clear()
    from app.main import app
    return TestClient(app)


def upload(test_client: TestClient) -> str:
    response = test_client.post("/api/documents/upload", files={"file": ("sample.docx", make_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert response.status_code == 201, response.text
    return response.json()["document_id"]


def test_health(tmp_path, monkeypatch):
    assert client(tmp_path, monkeypatch).get("/api/health").json()["status"] == "ok"


def test_upload_parse_export_and_targeted_edit(tmp_path, monkeypatch):
    test_client = client(tmp_path, monkeypatch)
    document_id = upload(test_client)
    parsed = test_client.get(f"/api/documents/{document_id}")
    assert parsed.status_code == 200
    paragraph = next(item for item in parsed.json()["segments"] if item["text"] == "Nama saya HUSNUL IBAT.")
    applied = test_client.post(f"/api/documents/{document_id}/apply", json={"userPrompt": "replace", "edits": [{"segmentId": paragraph["id"], "before": "HUSNUL IBAT", "after": "Lionel Messi", "action": "replace_text"}], "operations": []})
    assert applied.status_code == 200, applied.text
    exported = test_client.post(f"/api/documents/{document_id}/export")
    assert exported.status_code == 200
    assert Document(BytesIO(exported.content)).paragraphs[1].text == "Nama saya Lionel Messi."


def test_invalid_file_and_ai_configuration_error(tmp_path, monkeypatch):
    test_client = client(tmp_path, monkeypatch)
    invalid = test_client.post("/api/documents/upload", files={"file": ("bad.txt", b"no", "text/plain")})
    assert invalid.status_code == 415
    document_id = upload(test_client)
    ai = test_client.post(f"/api/documents/{document_id}/edit", json={"userPrompt": "Ubah nama", "segments": []})
    assert ai.status_code == 503
    assert ai.json()["detail"]["code"] == "ai_unavailable"


def test_cors_preflight(tmp_path, monkeypatch):
    response = client(tmp_path, monkeypatch).options("/api/documents/upload", headers={"Origin": "http://localhost:5173", "Access-Control-Request-Method": "POST"})
    assert response.status_code == 200
    assert response.headers["access-control-allow-origin"] == "http://localhost:5173"


def test_explicit_selected_text_is_planned_without_waiting_for_provider(tmp_path, monkeypatch):
    test_client = client(tmp_path, monkeypatch)
    monkeypatch.setenv("GEMINI_API_KEY", "configured-for-fast-path")
    from app.config import get_settings
    get_settings.cache_clear()
    document_id = upload(test_client)
    document = test_client.get(f"/api/documents/{document_id}").json()
    segment = next(item for item in document["segments"] if "HUSNUL IBAT" in item["text"])
    response = test_client.post(f"/api/documents/{document_id}/edit", json={
        "userPrompt": 'Edit this text: "HUSNUL IBAT" ganti dengan Lionel Messi',
        "segments": [segment],
    })
    assert response.status_code == 200, response.text
    assert response.json()["edits"][0]["after"] == "Lionel Messi"
