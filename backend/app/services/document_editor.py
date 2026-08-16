from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt

from app.models.documents import Operation, TextEdit
from app.utils.errors import api_error

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocumentEditor:
    def apply(
        self, path: Path, edits: list[TextEdit], operations: list[Operation]
    ) -> None:
        document = Document(path)
        paragraphs = document.paragraphs

        for edit in edits:
            self._apply_text_edit(paragraphs, document, edit)

        for operation in operations:
            self._apply_operation(document, paragraphs, operation)

        document.save(path)

    def _apply_text_edit(
        self, paragraphs: list, document: Document, edit: TextEdit
    ) -> None:
        if edit.action == "replace_image_with_uploaded":
            return

        if edit.action == "replace_image_with_text":
            image_index = self._image_index(edit.segmentId)
            shapes = document.inline_shapes
            if image_index < len(shapes):
                shape = shapes[image_index]
                drawing = shape._inline.getparent()
                if drawing is not None:
                    run = drawing.getparent()
                    if run is not None:
                        run.text = edit.after
                        run.remove(drawing)
            return

        if edit.segmentId.startswith("word/document.xml:table:"):
            self._apply_table_cell_edit(document, edit)
            return

        target = self._paragraph(paragraphs, edit.segmentId)

        if edit.action == "delete_text":
            for run in target.runs:
                run.text = ""
            return

        if edit.before not in target.text:
            raise api_error(
                422,
                "edit_target_mismatch",
                "A text edit no longer matches its target.",
                {"segmentId": edit.segmentId},
            )

        if len(target.runs) == 1:
            target.runs[0].text = target.runs[0].text.replace(
                edit.before, edit.after, 1
            )
        elif len(target.runs) > 1:
            self._replace_across_runs(target, edit.before, edit.after)
        else:
            target.text = target.text.replace(edit.before, edit.after, 1)

    @staticmethod
    def _replace_across_runs(paragraph, before: str, after: str) -> None:
        full_text = ""
        run_boundaries: list[tuple[int, int, object]] = []
        for run in paragraph.runs:
            start = len(full_text)
            full_text += run.text
            run_boundaries.append((start, len(full_text), run))

        idx = full_text.find(before)
        if idx == -1:
            return

        end_idx = idx + len(before)

        for start, end, run in run_boundaries:
            if end <= idx:
                continue
            if start >= end_idx:
                break

            run_start = max(idx, start) - start
            run_end = min(end_idx, end) - start

            if start <= idx < end:
                run.text = run.text[:run_start] + after + run.text[run_end:]
                after = ""
            else:
                run.text = run.text[:run_start] + run.text[run_end:]

    def _apply_table_cell_edit(self, document: Document, edit: TextEdit) -> None:
        if not edit.target:

            raise api_error(
                422,
                "invalid_operation",
                "A table edit requires a row/column target.",
                {"segmentId": edit.segmentId},
            )
        try:
            table_index = int(edit.segmentId.rsplit(":", 1)[1])
        except ValueError as exc:
            raise api_error(
                422, "invalid_operation", "Invalid table segmentId."
            ) from exc

        if table_index >= len(document.tables):
            raise api_error(
                422, "invalid_operation", "The requested table does not exist."
            )

        table = document.tables[table_index]
        row = edit.target.get("row", -1)
        column = edit.target.get("column", -1)
        if row < 0 or column < 0 or row >= len(table.rows):
            raise api_error(
                422, "invalid_operation", "The requested table cell is out of bounds."
            )
        if column >= len(table.rows[row].cells):
            raise api_error(
                422, "invalid_operation", "The requested table cell is out of bounds."
            )

        cell = table.rows[row].cells[column]
        cell_text = cell.text
        if edit.before not in cell_text:
            raise api_error(
                422,
                "edit_target_mismatch",
                "A table cell edit no longer matches its target.",
                {"segmentId": edit.segmentId, "row": row, "column": column},
            )

        full_text = "".join(p.text for p in cell.paragraphs)
        idx = full_text.find(edit.before)
        if idx == -1:
            return

        remaining = edit.after
        consumed = 0
        for paragraph in cell.paragraphs:
            para_text = paragraph.text
            para_start = consumed
            para_end = consumed + len(para_text)
            consumed = para_end

            if para_end <= idx:
                continue
            if para_start >= idx + len(edit.before):
                break

            overlap_start = max(idx, para_start)
            overlap_end = min(idx + len(edit.before), para_end)
            if overlap_start >= overlap_end:
                continue

            run_pos = 0
            for run in paragraph.runs:
                run_start = para_start + run_pos
                run_end = run_start + len(run.text)
                run_pos += len(run.text)

                if run_end <= overlap_start:
                    continue
                if run_start >= overlap_end:
                    break

                local_start = max(overlap_start, run_start) - run_start
                local_end = min(overlap_end, run_end) - run_start

                if run_start <= overlap_start < run_end:
                    run.text = run.text[:local_start] + remaining + run.text[local_end:]
                    remaining = ""
                else:
                    run.text = run.text[:local_start] + run.text[local_end:]

    def _apply_operation(
        self, document: Document, paragraphs: list, operation: Operation
    ) -> None:
        op_type = operation.type

        if op_type == "format_text":
            self._op_format_text(paragraphs, operation)
        elif op_type == "format_paragraph":
            self._op_format_paragraph(paragraphs, operation)
        elif op_type == "resize_image":
            self._op_resize_image(document, operation)
        elif op_type == "modify_page_layout":
            self._op_modify_page_layout(document, operation)
        elif op_type == "modify_heading_style":
            self._op_modify_heading_style(paragraphs, operation)
        elif op_type == "add_page_break":
            self._op_add_page_break(paragraphs, operation)
        elif op_type == "insert_text":
            self._op_insert_text(document, paragraphs, operation)
        elif op_type == "delete_text":
            self._op_delete_text(paragraphs, operation)
        elif op_type == "delete_image":
            self._op_delete_image(document, operation)
        else:
            raise api_error(
                422,
                "unsupported_operation",
                f"Operation '{op_type}' cannot be safely applied by the server yet.",
            )

    def _op_format_text(self, paragraphs: list, operation: Operation) -> None:

        paragraph = self._paragraph(paragraphs, operation.segmentId)
        props = operation.properties or {}
        for run in paragraph.runs:
            if "bold" in props:
                run.bold = bool(props["bold"])
            if "italic" in props:
                run.italic = bool(props["italic"])
            if "underline" in props:
                run.underline = bool(props["underline"])
            if "strikethrough" in props:
                run.font.strike = bool(props["strikethrough"])
            if "fontFamily" in props:
                run.font.name = str(props["fontFamily"])
            if "fontSize" in props:
                run.font.size = Pt(float(props["fontSize"]))
            if "color" in props:
                from docx.shared import RGBColor

                try:
                    run.font.color.rgb = RGBColor.from_string(
                        str(props["color"]).lstrip("#")
                    )
                except (ValueError, AttributeError):
                    pass
        if "alignment" in props and props["alignment"] in ALIGNMENT_MAP:
            paragraph.alignment = ALIGNMENT_MAP[props["alignment"]]

    def _op_format_paragraph(self, paragraphs: list, operation: Operation) -> None:

        paragraph = self._paragraph(paragraphs, operation.segmentId)
        props = operation.properties or {}
        pf = paragraph.paragraph_format

        if "spaceBefore" in props:
            pf.space_before = Pt(float(props["spaceBefore"]))
        if "spaceAfter" in props:
            pf.space_after = Pt(float(props["spaceAfter"]))
        if "lineSpacing" in props:
            pf.line_spacing = float(props["lineSpacing"])
        if "indentLeft" in props:
            pf.left_indent = Pt(float(props["indentLeft"]))
        if "indentRight" in props:
            pf.right_indent = Pt(float(props["indentRight"]))
        if "firstLine" in props:
            pf.first_line_indent = Pt(float(props["firstLine"]))
        if "alignment" in props and props["alignment"] in ALIGNMENT_MAP:
            paragraph.alignment = ALIGNMENT_MAP[props["alignment"]]

    def _op_resize_image(self, document: Document, operation: Operation) -> None:

        if operation.widthCm is None and operation.heightCm is None:
            raise api_error(
                422, "invalid_operation", "resize_image requires widthCm or heightCm."
            )

        image_index = self._image_index(operation.segmentId)
        shapes = document.inline_shapes
        if image_index >= len(shapes):
            raise api_error(
                422, "invalid_operation", "The requested image does not exist."
            )

        shape = shapes[image_index]
        original_ratio = shape.height / shape.width if shape.width else 1

        if operation.widthCm is not None and operation.heightCm is not None:
            shape.width = Cm(operation.widthCm)
            shape.height = Cm(operation.heightCm)
        elif operation.widthCm is not None:
            shape.width = Cm(operation.widthCm)
            shape.height = int(shape.width * original_ratio)
        else:
            shape.height = Cm(operation.heightCm)
            shape.width = (
                int(shape.height / original_ratio) if original_ratio else shape.width
            )

    def _op_modify_page_layout(self, document: Document, operation: Operation) -> None:

        props = operation.properties or {}
        for section in document.sections:
            if "marginTopCm" in props:
                section.top_margin = Cm(float(props["marginTopCm"]))
            if "marginBottomCm" in props:
                section.bottom_margin = Cm(float(props["marginBottomCm"]))
            if "marginLeftCm" in props:
                section.left_margin = Cm(float(props["marginLeftCm"]))
            if "marginRightCm" in props:
                section.right_margin = Cm(float(props["marginRightCm"]))
            if "pageSizeWidthCm" in props:
                section.page_width = Cm(float(props["pageSizeWidthCm"]))
            if "pageSizeHeightCm" in props:
                section.page_height = Cm(float(props["pageSizeHeightCm"]))
            if "orientation" in props:
                if props["orientation"] == "landscape":
                    section.orientation = WD_ORIENT.LANDSCAPE
                    if section.page_width < section.page_height:
                        section.page_width, section.page_height = (
                            section.page_height,
                            section.page_width,
                        )
                elif props["orientation"] == "portrait":
                    section.orientation = WD_ORIENT.PORTRAIT
                    if section.page_width > section.page_height:
                        section.page_width, section.page_height = (
                            section.page_height,
                            section.page_width,
                        )

    def _op_modify_heading_style(self, paragraphs: list, operation: Operation) -> None:
        paragraph = self._paragraph(paragraphs, operation.segmentId)
        props = operation.properties or {}

        if operation.level is not None:
            paragraph.style = f"Heading {operation.level}"

        for run in paragraph.runs:

            if "bold" in props:
                run.bold = bool(props["bold"])
            if "italic" in props:
                run.italic = bool(props["italic"])
            if "fontFamily" in props:
                run.font.name = str(props["fontFamily"])
            if "fontSize" in props:
                run.font.size = Pt(float(props["fontSize"]))

    def _op_add_page_break(self, paragraphs: list, operation: Operation) -> None:
        paragraph = self._paragraph(paragraphs, operation.segmentId)
        run = paragraph.add_run()
        run.add_break(docx.enum.text.WD_BREAK.PAGE)

    def _op_insert_text(
        self, document: Document, paragraphs: list, operation: Operation
    ) -> None:
        if not operation.text:
            raise api_error(
                422, "invalid_operation", "insert_text requires text content."
            )

        if operation.segmentId:
            ref_paragraph = self._paragraph(paragraphs, operation.segmentId)
            new_paragraph = OxmlElement("w:p")
            ref_paragraph._element.addnext(new_paragraph)
            from docx.text.paragraph import Paragraph

            inserted = Paragraph(new_paragraph, ref_paragraph._element.getparent())
            inserted.text = operation.text
        else:
            document.add_paragraph(operation.text)

    def _op_delete_text(self, paragraphs: list, operation: Operation) -> None:
        paragraph = self._paragraph(paragraphs, operation.segmentId)
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)

    def _op_delete_image(self, document: Document, operation: Operation) -> None:
        image_index = self._image_index(operation.segmentId)
        shapes = document.inline_shapes
        if image_index >= len(shapes):
            raise api_error(
                422, "invalid_operation", "The requested image does not exist."
            )

        shape = shapes[image_index]
        drawing = shape._inline.getparent()
        if drawing is not None:
            run = drawing.getparent()
            if run is not None:
                run.remove(drawing)

    @staticmethod
    def _paragraph(paragraphs: list, segment_id: str | None):

        if not segment_id or not segment_id.startswith("word/document.xml:paragraph:"):
            raise api_error(
                422, "invalid_operation", "This operation needs a paragraph segmentId."
            )
        try:
            index = int(segment_id.rsplit(":", 1)[1])
        except ValueError as exc:
            raise api_error(
                422, "invalid_operation", "Invalid paragraph segmentId."
            ) from exc
        if index >= len(paragraphs):
            raise api_error(
                422, "invalid_operation", "The requested paragraph does not exist."
            )
        return paragraphs[index]

    @staticmethod
    def _image_index(segment_id: str | None) -> int:
        if not segment_id or not segment_id.startswith("word/document.xml:image:"):
            raise api_error(
                422, "invalid_operation", "This operation needs an image segmentId."
            )
        try:
            return int(segment_id.rsplit(":", 1)[1])
        except ValueError as exc:
            raise api_error(
                422, "invalid_operation", "Invalid image segmentId."
            ) from exc
